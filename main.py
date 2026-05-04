import streamlit as st
import zipfile
import os
import shutil
import re
import json
import tempfile
import time
from pathlib import Path
from typing import List, Optional
import pdfplumber
from groq import Groq
from docx import Document
from openpyxl import load_workbook
import pandas as pd

# ------------------------------------------------------------
# Configuration
# ------------------------------------------------------------
EXTRACT_BASE = Path("./extracted_consultation")
EXTRACT_BASE.mkdir(exist_ok=True)

# ------------------------------------------------------------
# Universal text extraction from PDF, DOCX, Excel
# ------------------------------------------------------------
def extract_text_from_xls(file_path: Path) -> str:
    """Try multiple engines to read .xls files."""
    text = ""
    # Method 1: xlrd directly (requires xlrd==1.2.0)
    try:
        import xlrd
        wb = xlrd.open_workbook(str(file_path))
        for sheet in wb.sheets():
            text += f"\n--- Feuille: {sheet.name} ---\n"
            for row_idx in range(sheet.nrows):
                row_vals = sheet.row_values(row_idx)
                row_text = " | ".join(
                    str(cell).strip() if cell not in (None, "") else ""
                    for cell in row_vals
                )
                if row_text.strip():
                    text += row_text + "\n"
        if text.strip():
            return text
    except Exception as e1:
        pass

    # Method 2: pandas with xlrd engine
    try:
        all_sheets = pd.read_excel(str(file_path), sheet_name=None, engine="xlrd")
        for sheet_name, df in all_sheets.items():
            text += f"\n--- Feuille: {sheet_name} ---\n"
            text += df.fillna("").to_string(index=False) + "\n"
        if text.strip():
            return text
    except Exception as e2:
        pass

    # Method 3: pandas with openpyxl (sometimes .xls files are actually xlsx)
    try:
        all_sheets = pd.read_excel(str(file_path), sheet_name=None, engine="openpyxl")
        for sheet_name, df in all_sheets.items():
            text += f"\n--- Feuille: {sheet_name} ---\n"
            text += df.fillna("").to_string(index=False) + "\n"
        if text.strip():
            return text
    except Exception as e3:
        pass

    return ""


def extract_text_from_any_file(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    text = ""

    try:
        if suffix == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    # Also try to extract tables from PDF
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            row_text = " | ".join(str(cell) if cell else "" for cell in row)
                            if row_text.strip():
                                text += row_text + "\n"

        elif suffix in (".docx", ".doc"):
            if suffix == ".doc":
                st.warning(f"Fichier .doc non supporté directement : {file_path.name}")
                return ""
            doc = Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text += row_text + "\n"

        elif suffix == ".xlsx":
            wb = load_workbook(file_path, data_only=True)
            for sheet in wb.worksheets:
                text += f"\n--- Feuille: {sheet.title} ---\n"
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_text = " | ".join(
                            str(cell).strip() if cell is not None else ""
                            for cell in row
                        )
                        if row_text.strip(" |"):
                            text += row_text + "\n"

        elif suffix == ".xls":
            text = extract_text_from_xls(file_path)
            if not text:
                st.warning(f"⚠️ Impossible de lire le .xls: {file_path.name}. "
                           f"Assurez-vous que xlrd==1.2.0 est installé.")
                return ""

        else:
            st.warning(f"Format non supporté: {suffix} – {file_path.name}")
            return ""

    except Exception as e:
        st.error(f"Erreur d'extraction de {file_path.name}: {e}")
        return ""

    if not text.strip():
        st.warning(f"⚠️ Aucun texte extrait de {file_path.name} (fichier vide ou image scannée).")
    elif len(text.strip()) < 50:
        st.warning(f"⚠️ Texte très court dans {file_path.name}.")

    return text

# ------------------------------------------------------------
# ZIP & file handling
# ------------------------------------------------------------
def extract_zip(zip_path: Path, extract_to: Path, processed_paths: Optional[set] = None, depth: int = 0):
    if processed_paths is None:
        processed_paths = set()
    if depth > 10:
        st.warning(f"⚠️ Profondeur maximale atteinte pour {zip_path}")
        return

    abs_path = Path(zip_path).resolve()
    if abs_path in processed_paths:
        return
    processed_paths.add(abs_path)

    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(extract_to)
        st.write(f"📦 Extrait : `{zip_path}` → `{extract_to}`")

        # Recursively extract nested ZIP files
        for root, dirs, files in os.walk(extract_to):
            for file in files:
                if file.endswith('.zip'):
                    file_path = Path(os.path.join(root, file))
                    st.write(f"🔍 Sous-ZIP trouvé : `{file_path}`")
                    sub_extract_to = Path(os.path.join(extract_to, os.path.splitext(file)[0]))
                    extract_zip(file_path, sub_extract_to, processed_paths, depth + 1)

        # Delete all ZIP files after full recursive extraction (only at depth 0)
        if depth == 0:
            for root, dirs, files in os.walk(extract_to):
                for file in files:
                    if file.endswith(".zip"):
                        file_path = os.path.join(root, file)
                        os.remove(file_path)
                        st.write(f"🗑️ Supprimé : `{file_path}`")

    except zipfile.BadZipFile:
        st.error(f"❌ `{zip_path}` n'est pas un fichier ZIP valide ou est corrompu.")
    except Exception as e:
        st.error(f"❌ Une erreur est survenue : {e}")

def find_and_copy_files(src_dir: Path, dest_dir: Path, patterns: List[str]) -> List[Path]:
    dest_dir.mkdir(parents=True, exist_ok=True)
    copied = []
    for file in src_dir.rglob("*"):
        if file.is_file() and any(p in file.name.lower() for p in patterns):
            dest_path = dest_dir / file.name
            counter = 1
            while dest_path.exists():
                dest_path = dest_dir / f"{file.stem}_{counter}{file.suffix}"
                counter += 1
            shutil.copy2(file, dest_path)
            copied.append(dest_path)
    return copied

def move_files_by_lot_numbers(src_dir: Path, dest_dir: Path, lot_numbers: List[int]) -> List[Path]:
    dest_dir.mkdir(parents=True, exist_ok=True)
    moved = []
    regex = re.compile(r'(lot|n°|n[o°])\s*(\d+)', re.IGNORECASE)
    for file in src_dir.iterdir():
        if file.is_file():
            match = regex.search(file.name)
            if match and int(match.group(2)) in lot_numbers:
                dest_path = dest_dir / file.name
                counter = 1
                while dest_path.exists():
                    dest_path = dest_dir / f"{file.stem}_{counter}{file.suffix}"
                    counter += 1
                shutil.move(str(file), str(dest_path))
                moved.append(dest_path)
    return moved

# ------------------------------------------------------------
# Groq – lot detection
# ------------------------------------------------------------
def query_groq_for_lots(api_key: str, pdf_text: str, interests: List[str]) -> Optional[str]:
    max_input_chars = 8000
    if len(pdf_text) > max_input_chars:
        pdf_text = pdf_text[:max_input_chars] + "\n[...] (texte tronqué)"

    client = Groq(api_key=api_key)
    prompt = f"""
You are an expert in analyzing procurement and tender documents (RC files).

Analyze the following text and:
1. Generate a concise summary.
2. Identify ALL lots (Lot 1, Lot 2, etc.) in the document.
3. Detect lots relevant to these keywords: {', '.join(interests)}
4. For each relevant lot: lot number, description, matched keywords, reason.

Return STRICTLY valid JSON only (no markdown, no extra text):
{{
  "summary": "short summary",
  "total_lots_detected": number,
  "relevant_lots": [
    {{
      "lot_number": "Lot X",
      "description": "short description",
      "matched_keywords": ["kw1"],
      "reason": "why relevant"
    }}
  ],
  "relevant_lot_numbers": [4, 5, 7]
}}

TEXT:
{pdf_text}
"""
    try:
        resp = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0.2,
            max_tokens=4096,
        )
        return resp.choices[0].message.content
    except Exception as e:
        st.error(f"Groq API error (lot detection): {e}")
        return None

def extract_lot_numbers_from_groq_response(response_text: str) -> List[int]:
    lot_numbers = set()
    json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
    if json_match:
        try:
            data = json.loads(json_match.group())
            for item in data.get("relevant_lot_numbers", []):
                nums = re.findall(r'\d+', str(item))
                lot_numbers.update(int(n) for n in nums)
            for lot in data.get("relevant_lots", []):
                if "lot_number" in lot:
                    nums = re.findall(r'\d+', str(lot["lot_number"]))
                    lot_numbers.update(int(n) for n in nums)
        except json.JSONDecodeError:
            pass
    if not lot_numbers:
        for pattern in [r'lot\s*(?:n(?:°|o|uméro)?\s*)?(\d+)', r'"lot_number":\s*"?Lot\s*(\d+)"?']:
            matches = re.findall(pattern, response_text, re.IGNORECASE)
            lot_numbers.update(int(m) for m in matches)
    return sorted(lot_numbers)

# ------------------------------------------------------------
# Price estimation – improved prompt
# ------------------------------------------------------------
def calculate_price_for_file(api_key: str, file_path: Path) -> Optional[str]:
    file_text = extract_text_from_any_file(file_path)
    if not file_text or not file_text.strip():
        st.error(f"❌ Aucun texte extrait de {file_path.name}. Impossible d'estimer.")
        return None

    st.info(f"📄 Texte extrait de {file_path.name} : {len(file_text)} caractères")

    max_chars = 7000
    truncated = False
    if len(file_text) > max_chars:
        file_text = file_text[:max_chars] + "\n[...] (tronqué)"
        truncated = True

    client = Groq(api_key=api_key)
    prompt = f"""
Tu es un expert en estimation de coûts de construction en France.

Le texte ci-dessous est extrait d'un DPGF (Décomposition du Prix Global et Forfaitaire) ou CCTP.
Il contient des lignes de travaux avec codes, désignations, unités et parfois des quantités.

MISSION :
1. Identifie chaque poste de travaux (ligne avec code ou désignation).
2. Pour chaque poste :
   - Reprends la désignation exacte du document
   - Reprends l'unité si présente (U, ENS, ML, M², M³, PM, etc.)
   - Si la quantité est dans le document, utilise-la. Sinon, estime une quantité réaliste.
   - Attribue un prix unitaire HT réaliste selon les tarifs du marché français 2024.
3. Calcule : Total HT = Qté × Prix unitaire HT
4. En bas du tableau, affiche :
   - TOTAL HT = somme de tous les postes
   - TVA 20% = TOTAL HT × 0.20
   - TOTAL TTC = TOTAL HT + TVA

FORMAT DE SORTIE (tableau Markdown strict) :
| Code | Désignation | Qté | Unité | Prix Unit. HT (€) | Total HT (€) |
|------|-------------|-----|-------|-------------------|--------------|
| ...  | ...         | ... | ...   | ...               | ...          |

**TOTAL HT = X €**
**TVA 20% = Y €**
**TOTAL TTC = Z €**

RÈGLES :
- Ne retourne QUE le tableau et les totaux. Pas d'introduction ni d'explication.
- Si le texte contient des prix, utilise-les. Sinon estime selon le marché.
- Pour "ENS" ou "PM" → quantité = 1, estime le forfait global.
- Pour "U" → estime entre 1 et 20 selon le contexte.
- Sois précis et cohérent. N'invente pas de postes absents du texte.{"" if not truncated else chr(10) + "⚠️ Note: le texte a été tronqué. Traite uniquement les postes visibles."}

TEXTE EXTRAIT :
---
{file_text}
---
"""
    try:
        resp = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0.1,
            max_tokens=4096,
        )
        return resp.choices[0].message.content
    except Exception as e:
        st.error(f"Erreur Groq pour {file_path.name}: {e}")
        return None

# ------------------------------------------------------------
# Main app
# ------------------------------------------------------------
def main():
    st.set_page_config(page_title="Analyse Dossier de Consultation", layout="wide")
    st.title("📦 Analyse automatique d'un dossier de consultation")
    st.markdown(
        "Chargez votre ZIP → extraction → détection des lots pertinents via IA → "
        "estimation des prix (PDF, DOCX, XLSX, **XLS**)."
    )

    # ── Sidebar ──────────────────────────────────────────────
    with st.sidebar:
        st.header("🔑 Configuration Groq")
        groq_api_key = None
        if "GROQ_API_KEY" in st.secrets:
            groq_api_key = st.secrets["GROQ_API_KEY"]
            st.success("✅ Clé API depuis les secrets")
        else:
            groq_api_key = st.text_input("Clé API Groq", type="password",
                                         help="https://console.groq.com")
        st.markdown("---")
        st.header("📋 Mots-clés d'intérêt")
        default_interests = [
            "miroiterie", 
            "métallerie",
            "menuiserie extérieure",
             "Serrurerie",
             "Ascenseurs",
            "Escaliers mécaniques",
        ]
        edited = st.text_area("Un mot-clé par ligne",
                              value="\n".join(default_interests), height=300)
        interests = [l.strip() for l in edited.splitlines() if l.strip()]
        st.info(f"🧠 {len(interests)} mots-clés chargés")

    if not groq_api_key:
        st.warning("Veuillez entrer votre clé API Groq dans la barre latérale.")
        return

    # ── Upload ────────────────────────────────────────────────
    uploaded_zip = st.file_uploader(
        "📁 Téléchargez le fichier ZIP du dossier de consultation", type=["zip"]
    )
    if uploaded_zip is None:
        st.info("En attente du fichier ZIP…")
        return

    run_id      = str(int(time.time()))
    extract_root = EXTRACT_BASE / f"extract_{run_id}"
    finance_dir  = EXTRACT_BASE / f"finance_{run_id}"

    # ── Step 0 : extract ZIP ──────────────────────────────────
    with st.spinner("📦 Extraction du ZIP…"):
        temp_zip = extract_root.parent / uploaded_zip.name
        extract_root.mkdir(parents=True, exist_ok=True)
        with open(temp_zip, "wb") as f:
            f.write(uploaded_zip.getbuffer())
        extract_zip(temp_zip, extract_root)
        temp_zip.unlink()

        patterns = ["cctp", "dpgf", "rc"]
        copied   = find_and_copy_files(extract_root, finance_dir, patterns)

        # 1. Look for exact "rc" in filename
        rc_files = [f for f in finance_dir.iterdir() if re.search(r'\brc\b', f.name.lower())]
        
        # 2. Fallback: look for "rdc" in filename
        if not rc_files:
            rc_files = [f for f in finance_dir.iterdir() if "rdc" in f.name.lower()]
            if rc_files:
                st.info("ℹ️ Fichier RC trouvé via 'rdc'.")
        
        # 3. Fallback: look for "reglement" or "consultation" in filename
        if not rc_files:
            rc_files = [
                f for f in finance_dir.iterdir()
                if any(kw in f.name.lower() for kw in ["reglement", "règlement", "consultation"])
            ]
            if rc_files:
                st.info("ℹ️ Fichier RC trouvé via 'règlement de consultation'.")
        
        # 4. Nothing found
        if not rc_files:
            st.error("❌ Aucun RC trouvé dans le ZIP (testé : 'rc', 'rdc', 'règlement de consultation').")
            return

        rc_path = rc_files[0]
        st.success(f"✅ Fichier RC trouvé : {rc_path.name}")

        # Show all found files
        with st.expander("📂 Fichiers trouvés dans le dossier finance"):
            for f in finance_dir.iterdir():
                st.write(f"• {f.name}  ({f.suffix.upper()})")

    # ── Step 1 : Groq lot detection ───────────────────────────
    with st.spinner("🤖 Analyse du RC – détection des lots pertinents…"):
        rc_text = extract_text_from_any_file(rc_path)
        if not rc_text:
            st.error("Impossible d'extraire le texte du RC.")
            return
        groq_response = query_groq_for_lots(groq_api_key, rc_text, interests)
        if not groq_response:
            st.error("L'appel Groq a échoué pour la détection des lots.")
            return
        lot_numbers = extract_lot_numbers_from_groq_response(groq_response)

        # Show Groq summary
        try:
            data = json.loads(re.search(r'\{.*\}', groq_response, re.DOTALL).group())
            with st.expander("📋 Résumé du document RC"):
                st.write(data.get("summary", ""))
                st.write(f"**Lots totaux détectés :** {data.get('total_lots_detected', '?')}")
                for lot in data.get("relevant_lots", []):
                    st.markdown(
                        f"- **{lot.get('lot_number')}** – {lot.get('description')} "
                        f"_(mots-clés: {', '.join(lot.get('matched_keywords', []))})_"
                    )
        except Exception:
            pass

        if not lot_numbers:
            st.warning("Aucun lot détecté automatiquement.")
        else:
            st.success(f"📌 Lots pertinents : {lot_numbers}")

    # ── Step 2 : move matching files ──────────────────────────
    interesting_folder = finance_dir / "interesting_lot"
    moved_files: List[Path] = []

    if lot_numbers:
        moved_files = move_files_by_lot_numbers(finance_dir, interesting_folder, lot_numbers)
        if moved_files:
            st.success(f"✅ {len(moved_files)} fichier(s) déplacé(s) vers 'interesting_lot'.")
            for mf in moved_files:
                st.write(f"  • {mf.name}")
        else:
            st.warning("Aucun fichier CCTP/DPGF ne porte ces numéros de lot.")
    else:
        st.subheader("🔧 Saisie manuelle des lots")
        manual_input = st.text_input("Numéros de lots (séparés par des virgules)", value="")
        if manual_input:
            manual_lots = [int(n) for n in re.findall(r'\d+', manual_input)]
            if manual_lots:
                moved_files = move_files_by_lot_numbers(finance_dir, interesting_folder, manual_lots)
                if moved_files:
                    st.success(f"✅ {len(moved_files)} fichier(s) déplacé(s).")

    # ── Step 3 : price estimation ─────────────────────────────
    if moved_files:
        st.markdown("---")
        st.subheader("💰 Estimation des prix pour chaque lot")
        results = {}
        progress = st.progress(0)
        status   = st.empty()

        for idx, file_path in enumerate(moved_files):
            lot_match = re.search(r'(lot|n°|n[o°])\s*(\d+)', file_path.name, re.IGNORECASE)
            lot_label = lot_match.group(0).upper() if lot_match else file_path.stem
            status.text(f"⏳ Estimation pour {lot_label} ({file_path.suffix.upper()})…")

            price_result = calculate_price_for_file(groq_api_key, file_path)
            results[lot_label] = price_result or "❌ Échec de l'estimation."
            progress.progress((idx + 1) / len(moved_files))

        status.text("✅ Calculs terminés !")

        for lot, markdown_text in results.items():
            with st.expander(f"📊 Résultat pour {lot}", expanded=True):
                if markdown_text.startswith("❌"):
                    st.error(markdown_text)
                else:
                    st.markdown(markdown_text)

        # Combined download
        combined = "\n\n" + ("=" * 80 + "\n\n").join(
            f"# ESTIMATION POUR {lot}\n\n{text}" for lot, text in results.items()
        )
        st.download_button(
            "📥 Télécharger toutes les estimations (TXT)",
            data=combined,
            file_name="estimations.txt",
            mime="text/plain",
        )
    else:
        st.info(
            "Aucun fichier déplacé. Vérifiez que les noms des fichiers CCTP/DPGF "
            "contiennent le numéro de lot (ex: 'LOT 5', 'Lot 06', 'n°12')."
        )


if __name__ == "__main__":
    main()
